from __future__ import annotations

import json
import os
import subprocess
import sys
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]  # .../ml
EXAMPLES = ROOT / "examples"
DEFAULT_MODEL_DIR = ROOT / "artifacts" / "20251101T153949Z"


def read_json(path: Path) -> Any:
    return json.loads(path.read_text(encoding="utf-8"))


def run_module(
    module: str,
    args: List[str],
    payload: Optional[Any] = None,
    cwd: Optional[Path] = None,
    timeout_s: int = 60,
) -> Tuple[int, str, str]:
    cwd = cwd or ROOT
    stdin = ""
    if payload is not None:
        stdin = json.dumps(payload)

    proc = subprocess.run(
        [sys.executable, "-m", module, *args],
        input=stdin,
        text=True,
        capture_output=True,
        cwd=str(cwd),
        timeout=timeout_s,
        env={
            **os.environ,
            "PYTHONPATH": str(ROOT),
        },
    )
    return proc.returncode, proc.stdout.strip(), proc.stderr.strip()


def add_code_block(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)


def add_kv_table(doc: Document, rows: List[Tuple[str, str]]) -> None:
    table = doc.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Field"
    hdr_cells[1].text = "Meaning / behavior"

    for k, v in rows:
        cells = table.add_row().cells
        cells[0].text = k
        cells[1].text = v


def summarize_feature_columns(feature_columns: List[str]) -> Dict[str, List[str]]:
    groups: Dict[str, List[str]] = {
        "festival_*": [],
        "income": [],
        "request_*": [],
        "shipment_*": [],
        "batch_*": [],
        "ratios_and_deltas": [],
        "other": [],
    }

    for col in feature_columns:
        if col.startswith("festival_"):
            groups["festival_*"] .append(col)
        elif col.startswith("per_capita_income"):
            groups["income"].append(col)
        elif col.startswith("request_status_") or col in {"requested_kg", "unique_food_types", "request_count"}:
            groups["request_*"].append(col)
        elif col.startswith("incoming_") or col.startswith("outgoing_") or col == "avg_travel_time_minutes":
            groups["shipment_*"].append(col)
        elif col.startswith("produced_") or col.startswith("avg_batch_") or col == "avg_shelf_life_hours":
            groups["batch_*"].append(col)
        elif col in {
            "supply_demand_gap_kg",
            "net_flow_kg",
            "production_vs_demand_ratio",
            "request_to_supply_ratio",
        }:
            groups["ratios_and_deltas"].append(col)
        else:
            groups["other"].append(col)

    # remove empties
    return {k: v for k, v in groups.items() if v}


def main() -> None:
    model_dir = Path(os.environ.get("MODEL_DIR", str(DEFAULT_MODEL_DIR)))

    predict_records_path = EXAMPLES / "predict_records_example.json"
    predict_server_path = EXAMPLES / "predict_server_snapshot_example.json"
    transfer_path = EXAMPLES / "transfer_planner_example.json"

    metadata = read_json(model_dir / "metadata.json")
    feature_columns: List[str] = metadata.get("feature_summary", {}).get("feature_columns", [])

    # Run: inference (server snapshot)
    server_payload = read_json(predict_server_path)
    rc1, out1, err1 = run_module(
        "src.infer",
        ["--model-dir", str(model_dir)],
        payload=server_payload,
        cwd=ROOT,
        timeout_s=120,
    )

    # Run: inference (records)
    records_payload = read_json(predict_records_path)
    rc2, out2, err2 = run_module(
        "src.infer",
        ["--model-dir", str(model_dir)],
        payload=records_payload,
        cwd=ROOT,
        timeout_s=120,
    )

    # Run: transfer planner
    transfer_payload = read_json(transfer_path)
    rc3, out3, err3 = run_module(
        "src.transfer_planner",
        [
            "--mode",
            "all",
            "--max-pairs",
            "5",
            "--min-transfer-kg",
            "200",
            "--overstock-ratio",
            "0.8",
            "--understock-ratio",
            "0.4",
            "--target-ratio",
            "0.6",
        ],
        payload=transfer_payload,
        cwd=ROOT,
        timeout_s=120,
    )

    # Edge cases (captured, not destructive)
    edge_runs: List[Dict[str, str]] = []

    def record_edge_case(title: str, module: str, args: List[str], payload: Any = None) -> None:
        rc, out, err = run_module(module, args, payload=payload, cwd=ROOT, timeout_s=30)
        edge_runs.append(
            {
                "title": title,
                "command": f"python -m {module} {' '.join(args)}",
                "exit_code": str(rc),
                "stdout": out,
                "stderr": err,
            }
        )

    record_edge_case(
        "Edge 1: Empty payload (should error)",
        "src.infer",
        ["--model-dir", str(model_dir)],
        payload=None,
    )
    record_edge_case(
        "Edge 2: Wrong shape (neither records nor server snapshot)",
        "src.infer",
        ["--model-dir", str(model_dir)],
        payload={"hello": "world"},
    )
    record_edge_case(
        "Edge 3: Server snapshot yields 0 feature rows (missing dates)",
        "src.infer",
        ["--model-dir", str(model_dir)],
        payload={"freq": "M", "nodes": [], "requests": [], "shipments": [], "batches": []},
    )
    record_edge_case(
        "Edge 4: Model directory missing", 
        "src.infer",
        ["--model-dir", str(ROOT / "artifacts" / "__does_not_exist__")],
        payload=records_payload,
    )

    # Build DOCX
    doc = Document()

    # Title
    title = doc.add_paragraph("ML Model — Full Dry Run Lecture Notes (Runnable Examples)")
    title.runs[0].bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Generated: {datetime.now().isoformat(timespec='seconds')}")
    doc.add_paragraph(f"Model run directory used: {model_dir}")

    doc.add_heading("1) What this system is", level=1)
    doc.add_paragraph(
        "This ML subsystem is a Node.js gateway over a Python pipeline that runs unsupervised models. "
        "It does NOT directly plan shipments. It produces signals per region/time bucket:\n"
        "- KMeans cluster_id (which ‘type’ of region behavior a row resembles)\n"
        "- IsolationForest anomaly_score + is_anomaly (outlier detection)"
    )

    doc.add_heading("2) Two inference input formats", level=1)
    doc.add_paragraph(
        "The inference module (python -m src.infer) accepts either:\n"
        "A) { records: [...] } where each record is already a feature row\n"
        "B) Raw Server snapshot { nodes, requests, shipments, batches, freq } which is converted into feature rows on the fly."
    )

    doc.add_heading("2.1) Output schema", level=2)
    add_kv_table(
        doc,
        [
            ("count", "Number of result rows returned"),
            ("feature_columns", "The canonical numeric feature schema the model expects (from metadata.json)"),
            ("missing_feature_columns", "Features that were absent in the input and auto-filled with 0"),
            ("results[]", "Array of row predictions"),
            ("results[].cluster_id", "KMeans cluster assignment (integer)"),
            ("results[].anomaly_score", "IsolationForest decision_function score (higher = more normal, lower = more anomalous)"),
            ("results[].is_anomaly", "1 if IsolationForest predicts -1 (outlier), else 0"),
            ("results[].state/district/period_start", "Echoed key columns when present"),
        ],
    )

    doc.add_heading("3) Raw snapshot input attributes (what each field does)", level=1)
    doc.add_paragraph(
        "When you send raw Server data, Python runs feature engineering (aggregation) first. "
        "These are the key fields it expects and how it behaves when fields are missing."
    )

    doc.add_heading("3.1) nodes[]", level=2)
    add_kv_table(
        doc,
        [
            ("_id", "REQUIRED. Used as canonical node_mongo_id for joins"),
            ("state", "Used for grouping. Missing → filled as 'Unknown'"),
            ("district", "Used for grouping. Missing → filled as 'Unknown'"),
            ("type", "Not used for clustering directly, but used by transfer planner"),
            ("capacity_kg", "Used by transfer planner to compute utilization and capacity"),
            ("location.coordinates", "[lon, lat] used by transfer planner distance calculations"),
        ],
    )

    doc.add_heading("3.2) requests[]", level=2)
    add_kv_table(
        doc,
        [
            ("requestId", "Used for request_count and status counts"),
            ("requesterNode", "Join key to nodes._id to get state/district"),
            ("requiredBy_iso", "Converted to datetime and bucketed into period_start by freq"),
            ("items", "Can be list/dict/null. Normalized; sums required_kg and counts unique foodType"),
            ("items[].required_kg", "Summed into requested_kg"),
            ("items[].foodType", "Counted into unique_food_types"),
            ("status", "Pivoted into request_status_<status> counts"),
        ],
    )

    doc.add_heading("3.3) shipments[]", level=2)
    add_kv_table(
        doc,
        [
            ("shipmentId", "Counted into incoming_shipments/outgoing_shipments"),
            ("fromNode", "Join key to nodes._id for outgoing grouping"),
            ("toNode", "Join key to nodes._id for incoming grouping"),
            ("start_iso", "Datetime bucketed into period_start"),
            ("batchIds", "Can be list or single id; exploded for per-batch payload calculations"),
            ("travel_time_minutes", "Averaged into avg_travel_time_minutes (incoming)"),
        ],
    )

    doc.add_heading("3.4) batches[]", level=2)
    add_kv_table(
        doc,
        [
            ("_id", "Used as batch_mongo_id for shipment payload joining"),
            ("originNode", "Join key to nodes._id for produced_kg and produced_batches"),
            ("currentNode", "Used by transfer planner inventory maps"),
            ("manufacture_date", "Datetime bucketed into period_start for production features"),
            ("original_quantity_kg", "Renamed to initial_quantity_kg; used for produced_kg and shipment payload"),
            ("quantity_kg", "Renamed to current_quantity_kg; fallback shipment payload if initial missing"),
            ("freshnessPct", "Numeric; missing → default 100"),
            ("shelf_life_hours", "Numeric; missing → default 72"),
            ("status", "Used by transfer planner; only {stored,reserved} counted"),
        ],
    )

    doc.add_heading("4) Feature engineering: what gets computed", level=1)
    doc.add_paragraph(
        "All raw docs are rolled up into rows keyed by: state + district + period_start. "
        "These rollups create numeric columns that the ML models actually see."
    )

    doc.add_heading("4.1) Key columns", level=2)
    add_code_block(doc, "KEY_COLUMNS = ['state', 'district', 'period_start']")

    doc.add_heading("4.2) Derived feature columns (from the trained schema)", level=2)
    grouped = summarize_feature_columns(feature_columns)
    doc.add_paragraph(
        "The model’s canonical input schema comes from metadata.json and is the SAME schema used during inference. "
        "Incoming payloads can omit columns; missing columns are auto-filled with 0."
    )
    for group, cols in grouped.items():
        doc.add_paragraph(f"{group} ({len(cols)} columns):")
        add_code_block(doc, "\n".join(cols))

    doc.add_heading("4.3) Post-processed ratios/deltas", level=2)
    doc.add_paragraph(
        "After merging blocks, the pipeline computes these extra columns (with safe divide-by-zero handling):\n"
        "- supply_demand_gap_kg = incoming_kg - requested_kg\n"
        "- net_flow_kg = incoming_kg - outgoing_kg\n"
        "- production_vs_demand_ratio = produced_kg / requested_kg when requested_kg > 0 else 0\n"
        "- request_to_supply_ratio = requested_kg / incoming_kg when incoming_kg > 0 else 0"
    )

    doc.add_heading("5) The actual ML algorithms (what happens to the numbers)", level=1)
    doc.add_paragraph(
        "Both models operate on the numeric matrix built from feature_columns:\n"
        "- StandardScaler normalizes features\n"
        "- KMeans predicts cluster_id\n"
        "- IsolationForest produces anomaly_score (decision_function) and is_anomaly flag (predict == -1)."
    )

    doc.add_heading("6) Worked dry run: raw snapshot → engineered rows → predictions", level=1)
    doc.add_paragraph("Input used:")
    add_code_block(doc, json.dumps(server_payload, indent=2)[:12000])
    doc.add_paragraph("Python module executed:")
    add_code_block(doc, f"python -m src.infer --model-dir {model_dir}")
    doc.add_paragraph(f"Exit code: {rc1}")
    if err1:
        doc.add_paragraph("Stderr / warnings (may include sklearn version warnings):")
        add_code_block(doc, err1[:8000])
    doc.add_paragraph("Output JSON:")
    add_code_block(doc, out1[:12000])

    doc.add_heading("7) Worked dry run: records[] → predictions", level=1)
    doc.add_paragraph("Input used:")
    add_code_block(doc, json.dumps(records_payload, indent=2)[:6000])
    doc.add_paragraph("Output JSON:")
    add_code_block(doc, out2[:12000])

    doc.add_heading("8) Transfer planner (separate algorithm)", level=1)
    doc.add_paragraph(
        "This is NOT the clustering/anomaly model. It’s a deterministic planner that suggests:\n"
        "- warehouse_to_warehouse transfers when one warehouse is overstocked and another is understocked\n"
        "- farm_to_warehouse transfers when a farm has supply and a warehouse has capacity" 
    )
    doc.add_paragraph("Input used:")
    add_code_block(doc, json.dumps(transfer_payload, indent=2)[:6000])
    doc.add_paragraph("Output JSON:")
    add_code_block(doc, out3[:12000])

    doc.add_heading("9) Edge cases (what errors look like)", level=1)
    doc.add_paragraph(
        "These are intentionally-bad calls to show how the code fails and what message you get."
    )
    for edge in edge_runs:
        doc.add_heading(edge["title"], level=2)
        doc.add_paragraph(f"Command: {edge['command']}")
        doc.add_paragraph(f"Exit code: {edge['exit_code']}")
        if edge["stderr"]:
            doc.add_paragraph("stderr:")
            add_code_block(doc, edge["stderr"][:6000])
        if edge["stdout"]:
            doc.add_paragraph("stdout:")
            add_code_block(doc, edge["stdout"][:6000])

    out_path = EXAMPLES / "ML_Model_Lecture_Dry_Run.docx"
    doc.save(str(out_path))

    print(f"Wrote: {out_path}")


if __name__ == "__main__":
    main()
